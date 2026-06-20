import io
from typing import Optional
from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.database import get_db
from app.models.knowledge import KnowledgeEntry

router = APIRouter(prefix="/api/import", tags=["import"])


def parse_docx(file_bytes: bytes) -> list[dict[str, str]]:
    """解析 Word 文档，按段落拆分为知识条目"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))

    entries = []
    current_title = ""
    current_content: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 粗体或标题样式 → 作为新条目标题
        if para.style.name.startswith("Heading") or (para.runs and para.runs[0].bold):
            if current_title and current_content:
                entries.append({
                    "title_zh": current_title,
                    "content_zh": "\n".join(current_content),
                })
            current_title = text
            current_content = []
        else:
            current_content.append(text)

    # 最后一条
    if current_title and current_content:
        entries.append({"title_zh": current_title, "content_zh": "\n".join(current_content)})

    # 如果没识别到标题，整篇作为一个条目
    if not entries:
        all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if all_text:
            entries.append({"title_zh": "导入文档", "content_zh": all_text})

    return entries


def parse_pdf(file_bytes: bytes) -> list[dict[str, str]]:
    """解析 PDF，提取文本作为一个知识条目"""
    import fitz
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    full_text = "\n".join(text_parts).strip()

    if not full_text:
        return []

    # 尝试按空行分段落，第一段作为标题
    paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]
    if len(paragraphs) >= 2:
        return [{"title_zh": paragraphs[0][:100], "content_zh": full_text}]
    else:
        return [{"title_zh": "导入文档", "content_zh": full_text}]


def split_bilingual(content: str) -> tuple[str, str]:
    """
    尝试从内容中分离中俄双语。
    如果内容包含明显的俄语字符段（西里尔字母），
    则将内容按语言分成中/俄两部分。
    """
    has_cyrillic = any('\u0400' <= c <= '\u04FF' for c in content)

    if not has_cyrillic:
        return content, ""

    # 简单策略：中文部分 + 俄语部分（俄语在后面）
    lines = content.split("\n")
    zh_lines = []
    ru_lines = []
    current_lang = "zh"

    for line in lines:
        cyrillic_count = sum(1 for c in line if '\u0400' <= c <= '\u04FF')
        if cyrillic_count > len(line) * 0.3:
            current_lang = "ru"
        if current_lang == "ru":
            ru_lines.append(line)
        else:
            zh_lines.append(line)

    return "\n".join(zh_lines).strip(), "\n".join(ru_lines).strip()


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    module: str = Form("product"),
    db: AsyncSession = Depends(get_db),
):
    """上传 Word 或 PDF，自动解析入库"""
    if file.content_type is None:
        raise HTTPException(400, "无法识别文件类型")

    content = await file.read()
    filename = file.filename or ""

    # 解析
    if filename.lower().endswith(".docx"):
        entries = parse_docx(content)
    elif filename.lower().endswith(".pdf"):
        entries = parse_pdf(content)
    else:
        raise HTTPException(400, "仅支持 .docx 和 .pdf 格式")

    if not entries:
        raise HTTPException(400, "文档中未找到可导入的内容")

    # AI 自动分类
    from app.services.rag_service import call_llm

    async def auto_classify(title: str, cont: str) -> str:
        try:
            prompt = (
                "判断以下内容的模块，只回复一个英文词：\n"
                "product/sop/training/store/marketing/brand/franchise/operations/equipment/maintenance\n\n"
                f"标题：{title}\n内容：{cont[:200]}\n\n模块："
            )
            r = (await call_llm([{"role": "user", "content": prompt}])).strip().lower()
            valid = ["product","sop","training","store","marketing","brand","franchise","operations","equipment","maintenance"]
            if r in valid:
                return r
        except:
            pass
        return "product"

    # 写入数据库
    created = []
    for entry in entries:
        zh, ru = split_bilingual(entry["content_zh"])
        entry_module = module if module != "auto" else await auto_classify(entry["title_zh"], zh or entry["content_zh"])
        knowledge = KnowledgeEntry(
            module=entry_module,
            title_zh=entry["title_zh"],
            title_ru="",
            content_zh=zh or entry["content_zh"],
            content_ru=ru,
        )
        db.add(knowledge)
        created.append({"title": entry["title_zh"], "chars_zh": len(zh), "chars_ru": len(ru)})

    await db.commit()

    return {
        "message": f"成功导入 {len(created)} 条知识",
        "module": entry_module,
        "items": created,
    }


class PasteRequest(BaseModel):
    title_zh: str = ""
    content_zh: str
    content_ru: str = ""


@router.post("/paste")
async def paste_text(req: PasteRequest, db: AsyncSession = Depends(get_db)):
    """粘贴文本，AI 自动分类模块"""
    from app.services.rag_service import call_llm

    # AI 自动分类
    prompt = (
        f"以下是一段奶茶连锁店的知识内容。请判断它属于哪个模块，只回复模块英文名（一个词）。\n"
        f"可选模块：product(产品), sop(操作SOP), training(培训), store(门店), marketing(营销), "
        f"brand(品牌), franchise(特许), operations(运营), equipment(设备), maintenance(维修)\n\n"
        f"标题：{req.title_zh}\n内容：{req.content_zh[:300]}\n\n模块："
    )
    try:
        module = (await call_llm([{"role": "user", "content": prompt}])).strip().lower()
        if module not in ["product", "sop", "training", "store", "marketing", "brand", "franchise", "operations", "equipment", "maintenance"]:
            module = "product"
    except:
        module = "product"

    zh, ru = split_bilingual(req.content_zh)
    knowledge = KnowledgeEntry(
        module=module,
        title_zh=req.title_zh or req.content_zh[:40],
        title_ru="",
        content_zh=zh or req.content_zh,
        content_ru=ru or req.content_ru,
    )
    db.add(knowledge)
    await db.commit()

    return {"id": knowledge.id, "module": module, "message": "已添加"}
